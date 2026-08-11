from __future__ import annotations

import hashlib
import json
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


SOURCE = Path(r"C:\Users\Playdata\Downloads\모델링 및 평가_테스트 계획 및 결과 보고서_2팀.docx")
OUTPUT = Path(r"C:\Users\Playdata\Desktop\feature_backend\artifacts\docs\model_eval_report\template-evidence.json")


def emu_inches(value):
    return None if value is None else round(value / 914400, 4)


def twips(value):
    return None if value is None else round(value / 635, 2)


def color_value(color):
    if color is None or color.rgb is None:
        return None
    return str(color.rgb)


def run_info(run):
    fonts = run._element.get_or_add_rPr().rFonts
    return {
        "text": run.text,
        "font_name": run.font.name,
        "font_ascii": fonts.get(qn("w:ascii")) if fonts is not None else None,
        "font_east_asia": fonts.get(qn("w:eastAsia")) if fonts is not None else None,
        "size_pt": run.font.size.pt if run.font.size else None,
        "bold": run.bold,
        "italic": run.italic,
        "color": color_value(run.font.color),
    }


def paragraph_info(p):
    fmt = p.paragraph_format
    return {
        "text": p.text,
        "style": p.style.name,
        "alignment": p.alignment,
        "space_before_pt": fmt.space_before.pt if fmt.space_before else None,
        "space_after_pt": fmt.space_after.pt if fmt.space_after else None,
        "line_spacing": fmt.line_spacing,
        "left_indent_pt": fmt.left_indent.pt if fmt.left_indent else None,
        "first_line_indent_pt": fmt.first_line_indent.pt if fmt.first_line_indent else None,
        "keep_with_next": fmt.keep_with_next,
        "page_break_before": fmt.page_break_before,
        "runs": [run_info(run) for run in p.runs],
    }


def xml_attr(element, child_name, attr_name):
    child = element.find(qn(child_name)) if element is not None else None
    return child.get(qn(attr_name)) if child is not None else None


doc = Document(SOURCE)
section = doc.sections[0]
evidence = {
    "reference": {
        "path": str(SOURCE),
        "sha256": hashlib.sha256(SOURCE.read_bytes()).hexdigest(),
        "page_count": 4,
        "section_count": len(doc.sections),
    },
    "section": {
        "width_in": emu_inches(section.page_width),
        "height_in": emu_inches(section.page_height),
        "top_margin_in": emu_inches(section.top_margin),
        "bottom_margin_in": emu_inches(section.bottom_margin),
        "left_margin_in": emu_inches(section.left_margin),
        "right_margin_in": emu_inches(section.right_margin),
        "header_distance_in": emu_inches(section.header_distance),
        "footer_distance_in": emu_inches(section.footer_distance),
    },
    "paragraph_samples": {
        str(index): paragraph_info(doc.paragraphs[index])
        for index in [0, 1, 3, 4, 12, 16, 23, 25, 26, 28, 30, 32, 34, 37]
    },
    "tables": [],
}

for table_index, table in enumerate(doc.tables):
    tbl_pr = table._tbl.tblPr
    tbl_grid = table._tbl.tblGrid
    table_data = {
        "index": table_index,
        "rows": len(table.rows),
        "cols": len(table.columns),
        "style": table.style.name if table.style else None,
        "tblW": xml_attr(tbl_pr, "w:tblW", "w:w"),
        "tblW_type": xml_attr(tbl_pr, "w:tblW", "w:type"),
        "tblInd": xml_attr(tbl_pr, "w:tblInd", "w:w"),
        "grid_twips": [int(col.get(qn("w:w"))) for col in tbl_grid.gridCol_lst],
        "column_widths_in": [emu_inches(col.width) for col in table.columns],
        "header": [cell.text for cell in table.rows[0].cells],
        "header_paragraphs": [paragraph_info(cell.paragraphs[0]) for cell in table.rows[0].cells],
        "body_paragraphs": [paragraph_info(cell.paragraphs[0]) for cell in table.rows[1].cells],
    }
    evidence["tables"].append(table_data)

with ZipFile(SOURCE) as archive:
    evidence["package_parts"] = [
        {
            "path": item.filename,
            "size": item.file_size,
            "sha256": hashlib.sha256(archive.read(item.filename)).hexdigest(),
        }
        for item in archive.infolist()
    ]

OUTPUT.write_text(json.dumps(evidence, ensure_ascii=False, indent=2), encoding="utf-8")
print(OUTPUT)
