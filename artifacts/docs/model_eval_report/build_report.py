from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


SOURCE = Path(r"C:\Users\Playdata\Downloads\모델링 및 평가_테스트 계획 및 결과 보고서_2팀.docx")
OUTPUT = Path(
    r"C:\Users\Playdata\Desktop\feature_backend\artifacts\docs\모델링 및 평가_테스트 계획 및 결과 보고서_2팀_기능명세기반_재작성.docx"
)


def clear_runs(paragraph: Paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def set_paragraph_text(paragraph: Paragraph, text: str, *, size: float | None = None, bold: bool | None = None) -> None:
    clear_runs(paragraph)
    run = paragraph.add_run(text)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_labeled_paragraph(paragraph: Paragraph, label: str, text: str) -> None:
    clear_runs(paragraph)
    label_run = paragraph.add_run(label)
    label_run.bold = True
    paragraph.add_run(text)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.12


def clone_paragraph_before(target: Paragraph, source: Paragraph, text: str) -> Paragraph:
    element = deepcopy(source._p)
    target._p.addprevious(element)
    paragraph = Paragraph(element, target._parent)
    set_paragraph_text(paragraph, text)
    return paragraph


def clone_paragraph_after(anchor: Paragraph, source: Paragraph, text: str = "") -> Paragraph:
    element = deepcopy(source._p)
    anchor._p.addnext(element)
    paragraph = Paragraph(element, anchor._parent)
    set_paragraph_text(paragraph, text)
    return paragraph


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_cell_text(cell, text: str, *, size: float = 8.5, bold: bool = False, center: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    clear_runs(paragraph)
    run = paragraph.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def fill_table(table: Table, rows: list[list[str]], *, center_columns: set[int]) -> None:
    required_rows = len(rows) + 1
    while len(table.rows) < required_rows:
        table.add_row()
    while len(table.rows) > required_rows:
        table._tbl.remove(table.rows[-1]._tr)
    for row_index, values in enumerate(rows, start=1):
        row = table.rows[row_index]
        prevent_row_split(row)
        for column_index, value in enumerate(values):
            set_cell_text(
                row.cells[column_index],
                value,
                center=column_index in center_columns,
            )


shutil.copy2(SOURCE, OUTPUT)
doc = Document(OUTPUT)

# Capture source components before inserting new blocks.
meta_table = doc.tables[0]
functional_table = doc.tables[1]
performance_table = doc.tables[2]
stability_table = doc.tables[3]
history_table = doc.tables[4]
paragraphs = doc.paragraphs
source_bullet = paragraphs[24]
heading_34 = paragraphs[23]

# Cover metadata.
set_cell_text(meta_table.cell(2, 1), "2026-08-10", size=9.5)
set_cell_text(
    meta_table.cell(3, 1),
    "https://github.com/SKNETWORKS-FAMILY-AICAMP/SKN28-FINAL-2Team",
    size=9.0,
)
set_cell_text(meta_table.cell(4, 1), "SKN28 FINAL 2Team (2팀)", size=9.5)

# Overview and purpose.
set_paragraph_text(paragraphs[4], "테스트 기간: 2026-08-10 (병합 커밋 f123027 기준)")
set_paragraph_text(
    paragraphs[5],
    "목표: 탐나플랜의 화면·API·데이터 모델·LLM 일정 엔진에서 기능 명세를 추출하고, 실제 사용자 여정을 기준으로 기능·AI 품질·성능·예외 처리를 자동 검증한다.",
)
set_paragraph_text(
    paragraphs[7],
    "기능성 검증: 홈·패키지·인증·프로필·대화형 조건 입력·AI 일정 생성·동선·공유·슬롯 수정·찜·장바구니·예약·이력·평가 화면을 사용자 시나리오로 확인한다.",
)
set_paragraph_text(
    paragraphs[8],
    "AI 품질 검증: 5개 골든 질의를 10회 반복한 50케이스로 필수 장소 재현율, 콘텐츠 유형 준수, MySQL 근거성, 중복, 유사도 정렬을 평가한다.",
)
set_paragraph_text(
    paragraphs[9],
    "성능 검증: 실제 gpt-5-mini 일정 생성·대화 수정·패키지 추천, RAG 검색 지연시간과 프런트엔드 프로덕션 빌드 시간을 측정한다.",
)
set_paragraph_text(
    paragraphs[10],
    "안정성 및 에러 검증: 비인증 접근, 필수값 누락, 허용 범위 밖 입력, 중복 요청과 중복 취소를 API 계약에 맞게 처리하는지 확인한다.",
)

# Environment.
set_paragraph_text(
    paragraphs[13],
    "테스트 환경: 로컬 Windows NT 10.0.26200, x64, 논리 프로세서 8개",
)
set_paragraph_text(paragraphs[14], "RunPod GPU Pod: 본 자체 테스트에서는 사용하지 않음")
set_paragraph_text(paragraphs[15], "AWS EC2 (예정): 본 자체 테스트에서는 사용하지 않음")
set_paragraph_text(
    paragraphs[17],
    "언어·프레임워크: Python 3.14.3, Django 6.0.7, FastAPI 0.141.1, React 18.3.1, Vite 5.4.21, Node.js 24.14.1",
)
set_paragraph_text(
    paragraphs[18],
    "데이터베이스: MySQL(계정·여행·관광 원본 데이터), ChromaDB 1.5.9(jeju_places 벡터 인덱스)",
)
set_paragraph_text(
    paragraphs[19],
    "AI/ML: OpenAI SDK 2.53.0, 일정 생성·수정 모델 gpt-5-mini, AIHub 유사 여행 패턴, TourAPI 기반 장소 RAG",
)
set_paragraph_text(
    paragraphs[20],
    "실행 환경: Django 8000, RAG 평가 FastAPI 8001, Vite 5173 / 로컬 .venv 및 frontend node_modules",
)
set_paragraph_text(
    paragraphs[21],
    "기능 테스트: 라이브 HTTP E2E 자동화, 인앱 브라우저 UI 자동화, unittest, Django system check, FastAPI TestClient",
)
set_paragraph_text(
    paragraphs[22],
    "AI 품질·성능 테스트: feature/backend 골든 케이스 5개 × 10회, 라이브 LLM 일정 생성·수정, 응답 지연 및 빌드 산출물 측정",
)

# Add the missing 3.3 section using source-derived heading/list components.
clone_paragraph_before(heading_34, heading_34, "3.3 개발 기능 및 검증 범위")
development_items = [
    "인증·계정: Google/Kakao OAuth, JWT, 마이페이지·찜·예약 이력 접근 제어",
    "대화형 여행 조건: 기간·동행·스타일·추가 요청 수집 및 세션 상태 유지",
    "일정 엔진: AIHub 유사 여행 패턴과 RAG 장소 후보를 결합한 일정 생성·부분 수정",
    "장소 RAG: Chroma 벡터 검색, 콘텐츠 유형 필터, MySQL 원본 정보 복원, Top-K 30·후보 풀 100",
    "일정 UI: 날짜별 슬롯 편집, 지도·동선 조회, 공유·PDF·예약 흐름",
    "패키지 기능: 추천 패키지 3개, 상세·북마크·장바구니·예약 연동",
    "평가 기능: FastAPI 평가 작업, 골든 케이스, JSON/Markdown 보고서, /evaluation 대시보드",
]
for item in development_items:
    clone_paragraph_before(heading_34, source_bullet, item)

set_paragraph_text(
    paragraphs[24],
    "도구: tests/e2e/run_live_app_automation.py, 인앱 브라우저 자동화, Python unittest, Django manage.py, FastAPI TestClient, Vite build, RAG Evaluation API",
)
heading_34.paragraph_format.keep_with_next = True

# Functional test cases.
functional_rows = [
    ["TC-UI-01", "홈·핵심 CTA", "홈 진입 후 제목과 무료 일정 만들기 확인", "랜딩·CTA 표시", "PASS", "브라우저"],
    ["TC-UI-02", "이용 방법", "/how-to-use 직접 진입", "사용 안내 화면 표시", "PASS", "브라우저"],
    ["TC-UI-03", "패키지 목록", "전체 패키지 첫 페이지·페이지 버튼 확인", "카드 9건·페이지네이션", "PASS", "브라우저"],
    ["TC-UI-04", "패키지 상세", "패키지 카드를 선택해 모달 열기", "상세·닫기·예약 버튼", "PASS", "브라우저"],
    ["TC-AUTH-01", "인증 접근 제어", "비로그인 일정 API와 /mypage 접근", "API 401·로그인 이동", "PASS", "UI+HTTP"],
    ["TC-AUTH-02", "로그인 진입", "Google·Kakao 버튼 및 전용 사용자 JWT", "버튼 표시·토큰 200", "PASS", "UI+HTTP"],
    ["TC-ACC-01", "내 정보 조회", "인증 사용자 /api/accounts/me/", "이메일·닉네임·선호 반환", "PASS", "라이브 API"],
    ["TC-ACC-02", "내 정보 수정", "닉네임·선호 스타일 PATCH", "변경값 영구 저장", "PASS", "라이브 API"],
    ["TC-CHAT-01", "여행 조건 입력", "동행자 선택 후 날짜 단계 진입", "요약 갱신·날짜 입력 표시", "PASS", "브라우저"],
    ["TC-ITI-01", "일정 필수값", "종료일을 누락해 일정 생성 요청", "HTTP 400", "PASS", "라이브 API"],
    ["TC-ITI-02", "AI 일정 생성", "친구·2박 3일·액티비티 조건 전송", "3일·슬롯 생성·201", "PASS\n3일·15개", "gpt-5-mini"],
    ["TC-ITI-03", "일정 저장·조회", "생성 ID로 목록·상세 재조회", "사용자 소유 일정과 DAY 유지", "PASS", "라이브 API"],
    ["TC-ITI-04", "DAY별 동선", "생성 일정 route API 조회", "지도 좌표 경유지 반환", "PASS\n15좌표", "라이브 API"],
    ["TC-ITI-05", "일정 공유", "공유 토큰 발급 후 비로그인 조회", "공개 일정 HTTP 200", "PASS", "라이브 API"],
    ["TC-ITI-06", "맞춤 패키지", "생성 일정의 top_k=3 추천 요청", "추천 3건 반환", "PASS\n3건", "라이브 API"],
    ["TC-ITI-07", "대화 슬롯 교체", "DAY 1의 1번과 3번 교체 요청", "두 슬롯 순서 변경", "FAIL\n변경 없음", "라이브 LLM"],
    ["TC-ITI-08", "슬롯 순서 저장", "동일 항목을 days PATCH로 교환", "교환 순서 영구 저장", "PASS", "라이브 API"],
    ["TC-ITI-09", "일정 삭제", "본인 일정 DELETE", "HTTP 204", "PASS", "라이브 API"],
    ["TC-BMK-01", "패키지 찜", "추가·중복 추가·목록·해제", "단일 찜 유지 후 삭제", "PASS", "라이브 API"],
    ["TC-CART-01", "장바구니", "담기 후 수량 2·인원 4·날짜 수정", "옵션 저장·총액 재계산", "PASS", "라이브 API"],
    ["TC-CART-02", "수량 검증", "허용 범위 밖 수량 10 전송", "HTTP 400", "PASS", "라이브 API"],
    ["TC-RES-01", "예약 생성", "장바구니 항목으로 예약·결제", "확정·금액·항목·카트 정리", "PASS", "라이브 API"],
    ["TC-RES-02", "예약 취소", "확정 예약 취소 후 재취소", "취소 200·재취소 400", "PASS", "라이브 API"],
    ["TC-HIS-01", "이용 기록", "AI 대화 시작 기록 생성·조회", "본인 목록에서 확인", "PASS", "라이브 API"],
    ["TC-EVAL-01", "RAG 평가 UI", "평가 대시보드 진입", "케이스·실행 버튼 표시", "PASS", "브라우저"],
]
fill_table(functional_table, functional_rows, center_columns={0, 4, 5})

# AI quality note and table.
quality_intro = paragraphs[29]
quality_intro.style = doc.styles["Normal"]
set_paragraph_text(
    quality_intro,
    "평가 범위: 5개 골든 질의 × 10회 = 50케이스, 총 400개 Top-8 검색 결과. 결정론적 검색·필터 지표를 자동 평가했다.",
    size=9.5,
)
quality_intro.paragraph_format.space_after = Pt(6)

quality_tbl_element = deepcopy(performance_table._tbl)
quality_intro._p.addnext(quality_tbl_element)
quality_table = Table(quality_tbl_element, performance_table._parent)
quality_headers = ["TC ID", "평가 항목", "입력·조건", "평가 기준", "실제 결과", "판정"]
for index, value in enumerate(quality_headers):
    set_cell_text(quality_table.rows[0].cells[index], value, size=8.5, bold=True, center=True)
quality_rows = [
    ["TC-A01", "필수 장소 재현율", "한라수목원·섭지코지 각 10회", "required recall = 1.0", "20/20 반복 적중", "PASS"],
    ["TC-A02", "유형 필터 정확도", "식당·숙소·관광지 필터 30회", "유형 위반 0건", "compliance 1.0", "PASS"],
    ["TC-A03", "최소 결과·실행", "50케이스 Top-8 검색", "실행 성공·최소 결과 충족", "두 지표 모두 1.0", "PASS"],
    ["TC-A04", "중복·원본 근거성", "50케이스 400개 결과", "중복 0·MySQL 근거 100%", "두 지표 모두 1.0", "PASS"],
    ["TC-A05", "유사도·정렬", "50케이스 유사도 점수", "누락·내림차순 위반 0", "두 지표 모두 1.0", "PASS"],
]
fill_table(quality_table, quality_rows, center_columns={0, 4, 5})
quality_note = doc.add_paragraph()
set_paragraph_text(
quality_note,
    "제한사항: 자동 점수 100%는 사전에 정의된 검색·필터 게이트를 모두 통과했다는 뜻이다. LLM Judge, 실제 사용자의 주관 만족도와 일정 전체 동선 품질을 100% 보장하지 않는다.",
    size=9.0,
)
quality_note.paragraph_format.space_before = Pt(4)
quality_note.paragraph_format.space_after = Pt(4)
quality_table._tbl.addnext(quality_note._p)

# Performance cases.
performance_rows = [
    ["TC-P01", "AI 일정 생성", "gpt-5-mini·RAG 2박 3일 라이브 생성", "180초 이하", "86.4초 / 15슬롯", "PASS"],
    ["TC-P02", "대화 일정 수정", "슬롯 교체 요청 응답시간", "60초 이하", "12.2초 / 의미 변경 실패", "성능 PASS"],
    ["TC-P03", "RAG 검색 지연", "50케이스 검색 시간", "각 20초 이하", "평균 416.5ms / P95 503.5ms / 최대 1,786.9ms", "PASS"],
    ["TC-P04", "프로덕션 빌드", "Vite 539 modules 번들", "30초 이하·청크 500kB 이하", "8.14초 / 최대 청크 918.57kB", "조건부"],
]
fill_table(performance_table, performance_rows, center_columns={0, 4, 5})

# Stability/error cases.
paragraphs[32].paragraph_format.keep_with_next = True
stability_rows = [
    ["TC-S01", "인증 차단", "비로그인 일정 API 요청", "예상 401 / 실제 401", "PASS"],
    ["TC-S02", "필수값 검증", "end_date 없이 일정 생성", "예상 400 / 실제 400", "PASS"],
    ["TC-S03", "범위 검증", "장바구니 quantity=10", "예상 400 / 실제 400", "PASS"],
    ["TC-S04", "중복 찜", "동일 패키지를 두 번 찜", "단일 행 유지 / 실제 1건", "PASS"],
    ["TC-S05", "중복 예약 취소", "취소된 예약 재취소", "예상 400 / 실제 400", "PASS"],
]
fill_table(stability_table, stability_rows, center_columns={0, 3, 4})
for cell in stability_table.rows[0].cells:
    cell.paragraphs[0].paragraph_format.keep_with_next = True

# History.
history_rows = [
    ["기능 명세 재분석", "f123027", "2026-08-10", "범위 확정", "라우트·UI·API·모델 기반 사용자 기능 추출"],
    ["라이브 E2E 자동화", "신규 테스트", "2026-08-10", "26 PASS / 1 FAIL", "전용 사용자 생성·HTTP 실행·자동 정리"],
    ["브라우저 UI 자동화", "로컬 5173", "2026-08-10", "8 PASS", "공개 화면·로그인 이동·조건 입력 검증"],
    ["RAG 반복 평가", "eval-151952", "2026-08-10", "50/50 PASS", "5개 골든 질의 10회 반복"],
    ["회귀·시스템 검사", "f123027", "2026-08-10", "39 PASS / 1 SKIP", "unittest·Django·FastAPI 검증"],
    ["프런트 빌드", "Vite 5.4.21", "2026-08-10", "빌드 PASS", "539 modules·대형 청크 경고 1건"],
]
fill_table(history_table, history_rows, center_columns={1, 2, 3})

# Conclusion.
paragraphs[37].paragraph_format.page_break_before = True
conclusion = paragraphs[38]
set_labeled_paragraph(conclusion, "종합 판정: ", "조건부 통과")
conclusion_2 = clone_paragraph_after(
    conclusion,
    conclusion,
    "",
)
set_labeled_paragraph(
    conclusion_2,
    "검증 결과: ",
    "라이브 기능 E2E 27건 중 26건 PASS·1건 FAIL, 브라우저 UI 8건 전부 PASS, 회귀·어댑터 39건 PASS·1건 SKIP, Django 검사와 프런트 프로덕션 빌드를 통과했다.",
)
conclusion_3 = clone_paragraph_after(conclusion_2, conclusion, "")
set_labeled_paragraph(
    conclusion_3,
    "AI/RAG 결과: ",
    "gpt-5-mini로 2박 3일·15슬롯 일정을 86.4초에 생성했고 동선 좌표 15건과 추천 패키지 3건을 반환했다. RAG 50케이스는 50/50 통과, 평균 416.5ms였다.",
)
conclusion_4 = clone_paragraph_after(conclusion_3, conclusion, "")
set_labeled_paragraph(
    conclusion_4,
    "확인된 이슈: ",
    "'DAY 1의 1번과 3번 교체' 대화 요청은 HTTP 200이지만 순서가 변하지 않았다. 현재 ConditionDelta에 순서 교체 연산이 없고 notes-only 요청은 빈 변경으로 처리된다. PATCH 교환은 성공하지만 UI 직접 순서 변경 기능은 없다. 최대 청크 918.57kB 경고도 남아 있다.",
)
conclusion_5 = clone_paragraph_after(conclusion_4, conclusion, "")
set_labeled_paragraph(
    conclusion_5,
    "후속 권고: ",
    "LLM 변경 스키마에 swap/move(day, from_sequence, to_sequence) 연산을 추가하고 적용 후 순서를 검증해야 한다. UI에도 드래그 또는 위·아래 이동을 제공한다. 실제 Google/Kakao OAuth, 인증 상태 PDF 다운로드는 별도 수동 E2E가 필요하며 프런트 코드 분할도 권고한다.",
)

# Apply compact, source-consistent table typography and header repetition.
for table in [functional_table, quality_table, performance_table, stability_table, history_table]:
    table.autofit = False
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    if header_tr_pr.find(qn("w:tblHeader")) is None:
        header_tr_pr.append(OxmlElement("w:tblHeader"))
    prevent_row_split(table.rows[0])

doc.save(OUTPUT)
print(OUTPUT)
